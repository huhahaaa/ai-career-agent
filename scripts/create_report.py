"""
Create a fresh, well-formatted course report .docx for the ai-career-agent project.
Follows the YNU School of Software template structure.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

OUTPUT_DIR = r"D:\PythonProject\ai-career-agent\docs"
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "2026秋-工程实训（高级）-AI求职Agent实训报告.docx")

doc = Document()

# ============================================================
# Page setup
# ============================================================
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# ============================================================
# Helper functions
# ============================================================
def add_centered_para(text, size=12, bold=False, font_name='宋体', space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p

def add_left_para(text, size=12, bold=False, font_name='宋体', indent=None, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p

def add_body_para(text, indent=0.74, size=12):
    """Add a body text paragraph with first-line indent."""
    return add_left_para(text, size=size, indent=indent, space_after=3)

def set_cell_text(cell, text, bold=False, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Clear a cell and set its text."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    para.alignment = alignment
    para.clear()
    run = para.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return run

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def make_table_borders(table):
    """Add borders to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

def add_heading_para(text, level=1):
    """Add a heading with proper formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.size = Pt(16)
        run.bold = True
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.bold = True
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0.74)
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    add_centered_para("", size=12)

add_centered_para("云南大学软件学院期末课程报告", size=22, bold=True, space_after=4)
add_centered_para("Final Course Report", size=14, bold=False, font_name='Times New Roman', space_after=2)
add_centered_para("School of Software, Yunnan University", size=12, font_name='Times New Roman', space_after=24)

# Score box
score_box = doc.add_table(rows=2, cols=2)
score_box.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_text(score_box.rows[0].cells[0], "个人成绩", bold=True, size=14)
set_cell_text(score_box.rows[0].cells[1], "", size=14)
set_cell_text(score_box.rows[1].cells[0], "学  期：  2026年秋季学期", size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
set_cell_text(score_box.rows[1].cells[1], "", size=11)
make_table_borders(score_box)

add_centered_para("", size=12, space_after=12)

# Course info
info_lines = [
    "课程名称:    数字媒体技术实训(高级)",
    "授课教师:    黄老师",
    "课    题:    基于AI Agent的智能求职辅助平台设计与实现",
    "姓    名:    ___________",
    "联系电话:    ___________",
    "电子邮件:    ___________",
    "报告提交时间：  2026   年   7   月   27   日",
]
for line in info_lines:
    add_left_para(line, size=14, space_after=8)

doc.add_page_break()

# ============================================================
# TEAM MEMBER TABLE
# ============================================================
add_heading_para("项目组成员及分工", level=1)

add_left_para("（每位成员贡献总和为100%）", size=10, space_after=4)
add_left_para("注：1、个人贡献总和为100%；", size=10, space_after=2)
add_left_para("     2、分工内容必须互异，要求尽量详细具体；", size=10, space_after=6)

# Table: Team members
member_table = doc.add_table(rows=6, cols=5)
member_table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers_m = ["序号", "学  号", "姓  名", "负责模块", "贡献百分比(%)"]
for i, h in enumerate(headers_m):
    set_cell_text(member_table.rows[0].cells[i], h, bold=True, size=10)
    set_cell_shading(member_table.rows[0].cells[i], "D9E2F3")

members = [
    ("1", "", "何畅", "项目骨架、认证权限、数据库、后端集成、接口联调、阶段发布", "25%"),
    ("2", "", "黄林", "岗位采集清洗、数据审核、测试数据、失败案例、数据质量文档", "20%"),
    ("3", "", "牛帅", "前端页面、接口联调、交互优化、测试记录、演示材料", "20%"),
    ("4", "", "", "Embedding向量库、岗位检索匹配、知识库接入、匹配理由", "20%"),
    ("5", "", "", "简历审核、模拟面试Agent、评分反馈、面试题库、报告草稿", "15%"),
]
for i, (seq, sid, name, module, pct) in enumerate(members):
    row = member_table.rows[i + 1]
    set_cell_text(row.cells[0], seq, size=10)
    set_cell_text(row.cells[1], sid, size=10)
    set_cell_text(row.cells[2], name, size=10)
    set_cell_text(row.cells[3], module, size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(row.cells[4], pct, size=10)

make_table_borders(member_table)

add_centered_para("", size=12, space_after=12)

# ============================================================
# GRADING RUBRIC (Group)
# ============================================================
add_heading_para("期末考核小组成绩评定表", level=1)
add_left_para("年级： 2023级   专业： 数字媒体技术   学号：           姓名：", size=10, space_after=8)

group_rubric = doc.add_table(rows=9, cols=7)
group_rubric.alignment = WD_TABLE_ALIGNMENT.CENTER

gr_headers = ["评价指标", "分值", "指标内涵及评分标准", "", "", "", "得分"]
for i, h in enumerate(gr_headers):
    set_cell_text(group_rubric.rows[0].cells[i], h, bold=True, size=9)
    set_cell_shading(group_rubric.rows[0].cells[i], "D9E2F3")

gr_data = [
    ("选题价值", "5", "数字媒体应用场景选择准确，有较好的实用价值5", "应用场景比较准确，有一定实用价值4", "应用及价值一般3-2", "缺乏明确应用价值1"),
    ("需求分析", "5", "对关键技术进行充分的分析与论证，需求尽量完整准确5", "对关键技术进行分析，有较好的需求描述4", "需求分析描述一般3-2", "未对关键技术和需求进行分析描述1"),
    ("技术路线的可行程度", "10", "技术路线可行，有创新点10-9", "技术路线可行，有依据8-6", "可行性和依据一般5-3", "可行性较低2-1"),
    ("专业知识应用水平", "20", "能对前沿/高级Agent/数据库/知识库等理论或方法进行详尽准确的论述和总结20-17", "能对上述理论或方法进行较准确的论述和总结17-15", "专业知识应用一般14-10", "专业知识应用较差9-1"),
    ("使用现代工具", "10", "能熟练使用Agent/CV/媒体处理/数据库/Git等现代工具进行设计和开发10-9", "能较为熟练地使用多种现代工具进行设计和开发8-6", "对现代工具的使用不熟练5-3", "不能使用现代工具进行设计和开发2-1"),
    ("报告撰写质量", "20", "报告非常规范、文字表达好、逻辑结构清晰、图表格式规范20-18", "报告比较规范、文字表达较好、逻辑结构较好、图表格式较为规范17-15", "报告效果一般、文字表达一般、逻辑结构一般、图表格式一般14-10", "报告不规范、文字表达不好、逻辑结构不清晰、图表格式较差9-1"),
    ("综合得分（满分70分）", "综合得分（满分70分）", "", "", "", ""),
    ("备注：", "备注：", "任课教师签字：                         日期：", "", "", ""),
]

for i, row_data in enumerate(gr_data):
    row = group_rubric.rows[i + 1]
    for j, val in enumerate(row_data):
        set_cell_text(row.cells[j], val, size=8, alignment=WD_ALIGN_PARAGRAPH.LEFT if j >= 2 else WD_ALIGN_PARAGRAPH.CENTER)
    # Merge indicator content columns for first 6 rows
    if i < 6:
        row.cells[2].merge(row.cells[3])
        row.cells[2].merge(row.cells[4])
        row.cells[2].merge(row.cells[5])

make_table_borders(group_rubric)

doc.add_page_break()

# ============================================================
# INDIVIDUAL GRADING TABLE (5 members, same rubric)
# ============================================================
for member_idx in range(5):
    add_heading_para("期末考核个人成绩评定表", level=1)
    add_left_para(f"年级： 2023级   专业： 数字媒体技术   学号：             姓名：", size=10, space_after=8)

    ind_rubric = doc.add_table(rows=6, cols=7)
    ind_rubric.alignment = WD_TABLE_ALIGNMENT.CENTER

    ir_headers = ["评价指标", "分值", "指标内涵及评分标准", "", "", "", "得分"]
    for i, h in enumerate(ir_headers):
        set_cell_text(ind_rubric.rows[0].cells[i], h, bold=True, size=9)
        set_cell_shading(ind_rubric.rows[0].cells[i], "D9E2F3")

    ir_data = [
        ("小组成绩", "70", "见期末考核小组成绩评定表", "", "", ""),
        ("个人工作Git记录", "16",
         "承担核心模块，提交记录详实规范，有分支/合并/代码/文档/测试完整记录16-14",
         "完成主要模块，提交记录较充分，能说明个人贡献13-10",
         "完成部分任务，提交较少或记录零散，个人贡献说明一般9-4",
         "实际贡献很少，提交记录异常，无法说明个人贡献3-1"),
        ("协作与团队精神", "4",
         "主动沟通、积极分工协作，能推动团队进展4",
         "配合度良好，能按时完成分工和小组工作3",
         "沟通协作较少，配合度一般2",
         "不配合小组安排，影响项目进度1"),
        ("自主学习与解决问题", "10",
         "能独立学习Agent/前后端/数据库/数据处理等新技术并定位问题且改进10-9",
         "能学习并适应新技术，基本能独立解决问题8-7",
         "需要较多协助才能学习新技术，解决问题能力一般6-3",
         "缺乏自主学习能力，不能独立完成分配任务2-1"),
        ("综合得分（满分100分）", "综合得分（满分100分）", "", "", "", ""),
    ]

    for i, row_data in enumerate(ir_data):
        row = ind_rubric.rows[i + 1]
        for j, val in enumerate(row_data):
            set_cell_text(row.cells[j], val, size=8, alignment=WD_ALIGN_PARAGRAPH.LEFT if j >= 2 else WD_ALIGN_PARAGRAPH.CENTER)
        # Merge content columns for indicator rows with 4 levels
        if i in [0, 4]:
            row.cells[2].merge(row.cells[5])
        elif i in [1, 2, 3]:
            row.cells[2].merge(row.cells[3])
            row.cells[4].merge(row.cells[5])
            # For these rows, content is in cols 2-5, need to also merge
            # Actually the data already has 4 levels so cols 2,3,4,5 each have content
            pass

    make_table_borders(ind_rubric)

    # Add signature line
    add_left_para("备注：                                任课教师签字：                         日期：", size=9, space_after=4)

    if member_idx < 4:  # Don't page break after last one
        doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS placeholder
# ============================================================
doc.add_page_break()
add_heading_para("目  录", level=1)
add_centered_para("（请在Word中插入自动目录：引用 → 目录 → 自动目录）", size=10, space_after=12)

# ============================================================
# BODY CONTENT
# ============================================================
doc.add_page_break()

# ---- Section 1: 项目概述 ----
add_heading_para("1、项目概述", level=1)

add_body_para("1.1 项目背景")
add_body_para("在当前的就业市场环境下，高校毕业生面临着求职信息分散、简历优化困难、面试准备缺乏针对性等问题。传统的求职方式往往依赖学生个人在各大招聘平台手动搜索岗位、自行修改简历、独自练习面试，整个过程缺乏系统化的指导和智能化的辅助工具。")
add_body_para("随着大语言模型（LLM）、向量检索（Vector Search）和检索增强生成（RAG）等人工智能技术的快速发展，构建一个能够自动化处理岗位数据采集与审核、简历智能审核、岗位技能匹配、AI模拟面试和面试报告生成的一体化求职辅助平台成为可能。本项目'基于AI Agent的智能求职辅助平台'正是在这一背景下立项开发。")

add_body_para("1.2 项目目标")
add_body_para("本项目旨在为求职者（以数字媒体技术及相关专业学生为主要用户群体）提供一个完整的求职辅助工具链，具体目标包括：")
add_body_para("（1）岗位数据管理：支持从多个招聘渠道采集岗位数据，进行清洗、去重、技能提取和审核状态流转，构建结构化的岗位数据库。")
add_body_para("（2）简历智能审核：支持简历文件上传（PDF/Word/TXT），对简历进行结构化解析，识别优势与不足，给出修改建议。")
add_body_para("（3）岗位智能匹配：利用向量检索技术（Chroma + SentenceTransformer），根据简历内容和求职意向，从岗位库中检索最匹配的岗位，输出匹配分数、命中技能、缺失技能和能力缺口分析。")
add_body_para("（4）AI模拟面试：构建支持多种面试模式（技术面、HR面、压力面、反馈教练）的AI面试Agent，能够根据候选人简历和岗位要求个性化生成面试题，进行追问、评分和结构化反馈。")
add_body_para("（5）面试报告与看板：自动生成面试报告（含维度评分、STAR改写建议、练习计划），并通过数据看板展示求职全过程的统计图表。")

add_body_para("1.3 项目意义")
add_body_para("本项目的核心价值在于将AI技术实际应用于求职场景，帮助求职者更高效地定位岗位、优化简历、提升面试能力。同时，项目采用了Agent架构、向量检索、RAG知识库等前沿技术，为数字媒体技术专业的学生提供了一个综合性的工程实践平台，有助于理解和掌握现代AI应用系统的设计与开发方法。")

# ---- Section 2: 需求分析 ----
add_heading_para("2、需求分析", level=1)

add_body_para("2.1 功能需求")
add_body_para("根据课程要求和实际求职场景分析，系统需要满足以下功能需求：")
add_body_para("（1）用户认证与权限管理。系统需支持用户注册、登录、JWT Token鉴权，以及基于角色的权限控制（普通学生用户与审核员用户），确保不同角色只能访问对应权限的功能模块。")
add_body_para("（2）岗位数据管理。支持岗位数据的JSONL批量导入、单个创建、列表查询、状态流转（pending→approved/rejected），审核员可对岗位进行审核操作。已审核通过的岗位进入向量索引。")
add_body_para("（3）简历管理。支持PDF、Word、TXT格式简历文件上传，后端自动解析文本内容。简历支持版本管理，可查看历史版本和当前版本对比。")
add_body_para("（4）简历审核。对上传的简历进行规则+LLM双层审核，识别简历中的问题（空泛表达、夸大/绝对化表述、技能缺失等），生成审核报告。")
add_body_para("（5）岗位匹配。根据用户简历文本或求职意向，在向量库中检索相似岗位，返回匹配分数、匹配理由、命中技能和缺失技能。")
add_body_para("（6）AI面试Agent。支持技术面、HR面、压力面、反馈教练四种面试模式。系统根据候选人简历和目标岗位个性化生成8道面试题，支持追问机制，对回答进行5维度评分，生成STAR改写建议和练习计划。")
add_body_para("（7）面试报告。面试结束后自动生成结构化报告，包含总分、维度得分、逐题详情、STAR建议和练习计划。")
add_body_para("（8）数据看板。展示简历数、岗位数、面试次数、平均分、热门技能需求、岗位城市分布、最近面试记录等统计图表。")
add_body_para("（9）Agent日志。记录Agent每次关键调用的操作类型、状态、耗时、输入输出摘要，用于系统可观测性和调试。")

add_body_para("2.2 非功能需求")
add_body_para("（1）性能要求：API响应时间应在合理范围内（p95 < 5s），向量检索应在秒级完成。")
add_body_para("（2）可靠性要求：LLM离线时系统应能降级运行，通过规则引擎完成基本功能，不影响核心流程。")
add_body_para("（3）可用性要求：前端界面应直观易用，面试界面支持语音输入和TTS语音播报。")
add_body_para("（4）安全性要求：密码bcrypt加密存储、JWT token鉴权、敏感数据不落地到前端日志。")
add_body_para("（5）可扩展性要求：知识库材料（岗位画像、技能词典、面试题库）应独立于代码，支持热更新。")

add_body_para("2.3 用户角色分析")
add_body_para("系统定义两类用户角色：")
add_body_para("（1）学生用户（student角色）：可使用简历上传、简历审核、岗位匹配、模拟面试、面试历史、数据看板等功能。")
add_body_para("（2）审核员用户（reviewer角色）：除学生功能外，额外拥有岗位审核、岗位索引重建、系统指标查看等管理权限。")

# ---- Section 3: 系统设计 ----
add_heading_para("3、系统设计", level=1)

add_body_para("3.1 系统架构")
add_body_para("本系统采用前后端分离的B/S架构，整体分为三层：")
add_body_para("（1）前端展示层（React + Vite）：负责用户交互界面，包含登录注册、数据看板、简历管理、岗位管理、岗位匹配、模拟面试、面试历史等9个核心页面。通过RESTful API与后端通信。")
add_body_para("（2）后端服务层（Python + FastAPI）：提供业务逻辑处理，包括用户认证、岗位管理、简历审核、向量匹配、面试Agent、报告生成等功能模块。采用依赖注入模式管理数据库会话和用户认证。")
add_body_para("（3）数据存储层：使用SQLite作为主数据库（通过SQLAlchemy ORM访问），Chroma作为向量数据库（存储岗位嵌入向量），本地文件系统存储简历上传文件。")

add_body_para("3.2 技术架构")
add_body_para("系统核心技术栈如下：")
add_body_para("- 后端框架：Python 3.11 + FastAPI 0.110.3")
add_body_para("- 数据库：SQLite + SQLAlchemy 2.0.31 ORM")
add_body_para("- 向量检索：Chroma + SentenceTransformer（bge-large-zh-v1.5）")
add_body_para("- 认证：bcrypt 5.0.0 + python-jose JWT")
add_body_para("- LLM集成：OpenAI兼容接口（DeepSeek），支持LLM离线降级")
add_body_para("- 前端：React 18 + Vite + Recharts（图表库）")
add_body_para("- 简历解析：PyPDF + python-docx")
add_body_para("- 语音：浏览器Web Speech API（语音输入）+ Edge TTS / 豆包TTS（语音输出）")

add_body_para("3.3 数据库设计")
add_body_para("系统设计以下核心数据表：")
add_body_para("（1）users表：用户认证信息（username, email, hashed_password, role, is_active, created_at）。")
add_body_para("（2）job_postings表：岗位信息（source_id, category, title, company, location, salary_range, education, experience, responsibilities, requirements, skills（JSON）, publish_time, source_site, source_link, collected_at, status, audit_comment）。")
add_body_para("（3）resumes表：简历元信息（user_id, title, filename, file_path, source_type, current_version_number, is_default）。")
add_body_para("（4）resume_versions表：简历版本内容（resume_id, version_number, content, file_name, created_at）。")
add_body_para("（5）resume_audit_reports表：简历审核报告（resume_id, user_id, report_data（JSON）, created_at）。")
add_body_para("（6）matching_records表：匹配记录（user_id, job_id, resume_id, total_score, skill_scores, matched_skills, missing_skills, created_at）。")
add_body_para("（7）interview_sessions表：面试会话（user_id, job_id, interview_mode, status, score, report_data, started_at, finished_at）。")
add_body_para("（8）messages表：聊天消息（session_id, role, content, created_at），其中role=system的消息存储Agent内部状态。")
add_body_para("（9）agent_logs表：Agent调用日志（user_id, session_id, operation, status, duration_ms, input_summary, output_summary, error_message, created_at）。")

add_body_para("3.4 面试Agent设计")
add_body_para("面试Agent是本系统的核心模块，其设计如下：")
add_body_para("（1）Agent状态持久化：Agent状态（面试题列表、当前进度、用户回答、评分）序列化为JSON存储在messages表中role=system的记录里，通过_load_agent_state()和_save_agent_state()进行读写。")
add_body_para("（2）题目生成策略：采用「简历解析→岗位需求分析→LLM生题」三段式Pipeline。先用LLM解析简历提取技能和项目经历，再分析岗位JD生成面试关注点，最后LLM综合两方信息生成8道个性化题目。题库参考来自interview_question_bank.json中的72道题目（30道通用+42道岗位专属）。")
add_body_para("（3）追问决策：采用规则+LLM混合策略。规则层检查回答长度（<50字）、技术关键词覆盖率、量化数据有无、题目要点覆盖度；通过规则判断后由LLM生成个性化追问文本。")
add_body_para("（4）评分体系：5维度评分（content_relevance 25分、professional_accuracy 25分、clarity 20分、star_completeness 20分、position_match 10分，满分100分）。LLM评分后经过规则层兜底修正（空泛表达封顶clarity、夸大表达封顶professional_accuracy、漏答要点封顶content_relevance）。")
add_body_para("（5）LLM离线降级：所有LLM调用点都有fallback机制，LLM离线时系统可降级为纯规则运行，不影响面试流程完整性。")

add_body_para("3.5 知识库设计")
add_body_para("系统知识库由四层材料构成：")
add_body_para("（1）技能规范层：skill_dictionary.json定义20条技能同义词归一化规则（如Python→[python, Python开发, Python编程]），用于岗位数据清洗时统一技能名称。")
add_body_para("（2）岗位能力模型层：role_profiles.json定义6类岗位（前端开发、后端开发、产品经理、运营、算法/机器学习、数字媒体/内容）的能力画像，含must_have技能、preferred技能和evidence_signals。")
add_body_para("（3）面试题库层：interview_question_bank.json含30道通用题+42道岗位专属题，覆盖HR面/技术面/压力面/反馈教练四种模式；interview_questions.jsonl含30道带评分点和追问标签的结构化题目。")
add_body_para("（4）数据质量层：test_cases.json（10个测试用例）、day2_failure_cases.json（6个失败案例）、job_sources.csv（24条岗位来源追踪）、resume_samples.jsonl（10份脱敏简历样本）、job_jd_samples.jsonl（10条JD样本）。")

# ---- Section 4: 系统实现 ----
add_heading_para("4、系统实现", level=1)

add_body_para("4.1 开发环境与工具")
add_body_para("- 操作系统：Windows 11")
add_body_para("- 开发语言：Python 3.11（后端）、JavaScript/JSX（前端）")
add_body_para("- IDE：VS Code / PyCharm")
add_body_para("- 版本管理：Git，托管于GitHub")
add_body_para("- 包管理：pip（后端）、npm（前端）")
add_body_para("- 测试框架：pytest + httpx（后端）")
add_body_para("- LLM服务：DeepSeek API（deepseek-chat / deepseek-v4-flash）")
add_body_para("- 向量模型：SentenceTransformer bge-large-zh-v1.5")

add_body_para("4.2 后端核心实现")
add_body_para("4.2.1 FastAPI应用工厂模式")
add_body_para("后端采用FastAPI的lifespan模式管理应用生命周期：启动时初始化数据库表、检查向量存储可用性；关闭时清理资源。CORS中间件配置允许前端跨域访问。")
add_body_para("4.2.2 异常处理体系")
add_body_para("自定义AppException异常类（含status_code、code、message、data、headers），注册四个全局异常处理器（AppException、RequestValidationError、HTTPException、通用Exception），统一返回{code, message, data}格式的API响应。")
add_body_para("4.2.3 认证与鉴权")
add_body_para("采用bcrypt哈希存储密码（含72字节长度限制检查），JWT Bearer Token鉴权（含type: 'access'声明），通过get_current_user()依赖注入获取当前用户，require_roles(*allowed_roles)依赖工厂实现角色权限控制。")
add_body_para("4.2.4 向量存储")
add_body_para("VectorStore采用线程安全的双重检查锁定单例模式（Double-Checked Locking），使用lazy import避免未安装依赖时导入失败。支持upsert_job()单条索引和search()批量检索，返回余弦相似度分数。")
add_body_para("4.2.5 岗位匹配")
add_body_para("match_resume_to_jobs()服务函数将简历文本与求职意向拼接为查询字符串，调用VectorStore.search()检索相似岗位，结合技能覆盖度计算综合匹配分数。匹配结果包含命中技能（matched_skills）、缺失技能（missing_skills）和匹配分数。")
add_body_para("4.2.6 简历审核")
add_body_para("audit_resume_text()采用规则+LLM加权混合审计：规则层检查空泛词（VAGUE_PHRASES）、夸大词（BIASED_PHRASES）、量化数据有无；LLM层进行深度分析。最终分数由30%规则分数+70%LLM分数加权计算。")
add_body_para("4.2.7 面试Agent")
add_body_para("interview_agent.py为核心实现文件（约1200行），包含以下关键函数：")
add_body_para("- start_interview()：解析简历、分析岗位需求、生成8道个性化题目，初始化Agent状态。")
add_body_para("- evaluate_answer()：接收用户回答，判断是否需要追问或直接评分。")
add_body_para("- should_followup()：复合规则判断是否需要追问（长度<50字/无技术词/无量化/要点覆盖不足）。")
add_body_para("- _generate_followup()：LLM生成个性化追问文本。")
add_body_para("- _score_answer()：LLM评分+规则兜底修正+题目要点覆盖检查。")
add_body_para("- finish_interview()：计算维度平均分、生成STAR改写建议、练习计划和总评。")
add_body_para("Agent调用日志通过agent_operation_log()上下文管理器自动记录，包含操作类型、状态、耗时、输入摘要和输出摘要。")

add_body_para("4.3 前端核心实现")
add_body_para("4.3.1 页面结构")
add_body_para("前端包含9个核心页面：Dashboard（数据看板）、ResumeUpload（简历上传）、ResumeReview（简历审核）、JobCollection（岗位采集）、JobReview（岗位审核）、JobSearchMatch（岗位匹配搜索）、JobComparison（岗位对比）、MockInterview（模拟面试）、InterviewHistory（面试历史）。")
add_body_para("4.3.2 Mock API模式")
add_body_para("API客户端采用mockFallback模式：后端在线时请求真实API，后端离线时自动降级使用内嵌mock数据返回，保证前端开发不依赖后端状态。")
add_body_para("4.3.3 面试交互")
add_body_para("模拟面试页面集成了语音输入（浏览器Web Speech API）、摄像头画面采集和TTS语音播报（Edge TTS或豆包TTS），提供沉浸式面试体验。聊天气泡展示每轮得分、追问和反馈。")
add_body_para("4.3.4 图表可视化")
add_body_para("数据看板使用Recharts库实现6种图表：个人技能分布（水平柱状图）、热门技能需求（垂直柱状图）、个人能力vs岗位要求（雷达图）、多岗位匹配得分（水平柱状图）、面试得分趋势（折线图）、岗位城市分布（饼图）。")

add_body_para("4.4 数据工程")
add_body_para("4.4.1 岗位数据")
add_body_para("系统维护100条中文岗位数据（data/processed/jobs_chinese.jsonl），覆盖6大类岗位、46家公司、11个城市，每条包含19个字段（source_id至audit_comment）。通过脚本生成并标注薪资范围。")
add_body_para("4.4.2 数据质量保障")
add_body_para("建立10条数据质量测试用例（test_cases.json），覆盖缺字段、缺来源、内容过短、重复检测、相似检测、技能归一化、乱码、日期缺失、审核流转、状态排除等场景。6条系统失败案例（day2_failure_cases.json）用于测试异常输入处理。")
add_body_para("4.4.3 岗位审核流程")
add_body_para("岗位数据导入后默认status为'pending'，审核员可标记为'approved'（进入向量索引）或'rejected'（排除出匹配），审核意见写入audit_comment字段。")

# ---- Section 5: 系统测试 ----
add_heading_para("5、系统测试", level=1)

add_body_para("5.1 测试策略")
add_body_para("本系统采用多层次的测试策略：")
add_body_para("（1）单元测试：对核心服务函数（面试Agent评分逻辑、规则引擎、岗位清洗、JWT认证等）编写pytest单元测试。")
add_body_para("（2）集成测试：对API接口（认证、岗位CRUD、匹配、面试流程等）使用httpx编写集成测试。")
add_body_para("（3）前端构建验证：每次提交前执行npm run build确认前端可正常构建。")
add_body_para("（4）手动验收测试：对照Day2计划中的验收标准逐项检查功能完整性。")

add_body_para("5.2 后端测试")
add_body_para("后端测试使用pytest框架，测试文件位于backend/tests/目录。测试覆盖以下核心场景：认证测试（注册、登录、Token验证、权限控制、密码加密）、岗位管理测试（岗位创建、列表查询、状态审核流转、岗位索引）、岗位匹配测试（匹配接口、匹配历史、技能分数字段）、面试Agent测试（面试启动、回答评估、追问触发、报告生成）、Agent日志测试（日志写入、操作记录、状态记录）、数据看板测试（统计接口、数据正确性）。")
add_body_para("所有测试在LLM离线状态下使用mock数据进行，确保测试结果稳定可复现。")

add_body_para("5.3 核心功能验收")
add_body_para("根据Day2验收标准，逐项验证如下：")
add_body_para("（1）完整流程可跑通：注册登录→导入/审核岗位→更新向量索引→简历匹配→Agent面试→结束报告→查看历史。验证通过。")
add_body_para("（2）看板至少3个图表模块：实现个人技能分布、热门技能需求、个人能力vs岗位要求、多岗位匹配得分、面试趋势、城市分布共计6个图表。验证通过。")
add_body_para("（3）Agent至少3个工具节点：简历解析、岗位检索、面试出题、追问评分、STAR反馈共计5个节点。验证通过。")
add_body_para("（4）面试模式切换：技术面、HR面、压力面、反馈教练共计4种模式。验证通过。")
add_body_para("（5）评分结构化：每轮反馈包含得分、优点、问题、改进建议。验证通过。")
add_body_para("（6）LLM离线降级：LLM不可用时系统降级为纯规则评分，不影响面试流程完整性。验证通过。")

add_body_para("5.4 失败案例测试")
add_body_para("系统设计了6个失败测试案例，覆盖异常场景：FAIL-001（简历太短→触发追问）、FAIL-002（岗位无来源→不能approved）、FAIL-003（匹配为空→返回no_match）、FAIL-004（回答过短→触发追问+降分）、FAIL-005（含隐私信息→脱敏处理）、FAIL-006（岗位过期→标记为stale）。")

# ---- Section 6: 总结 ----
add_heading_para("6、总结", level=1)

add_body_para("6.1 项目成果")
add_body_para("本项目成功构建了一个基于AI Agent的智能求职辅助平台，实现了从岗位数据管理、简历审核、岗位匹配到AI模拟面试的完整闭环。主要成果包括：")
add_body_para("（1）后端系统：基于FastAPI构建了完整的RESTful API服务，包含用户认证与鉴权、岗位管理、简历管理、向量匹配、面试Agent、报告生成、数据看板、Agent日志等9大功能模块，代码总量约8000行。")
add_body_para("（2）前端系统：基于React+Vite构建了9个交互页面，包含数据看板、简历管理、岗位管理、岗位匹配、模拟面试等核心功能，支持语音输入和TTS语音播报。")
add_body_para("（3）数据资产：建立了包含100条中文岗位数据、6类岗位能力画像、20条技能同义词规则、72道面试题、10份脱敏简历样本、10条数据测试用例、6个失败案例的数据资产库。")
add_body_para("（4）AI Agent：实现了支持4种面试模式的智能面试Agent，具备题目个性化生成、智能追问、5维度评分、STAR改写建议和练习计划生成能力，且支持LLM离线降级运行。")
add_body_para("（5）知识库体系：构建了技能规范、岗位能力模型、面试题库、数据质量标准四层知识库，为RAG检索增强生成奠定基础。")

add_body_para("6.2 技术收获")
add_body_para("通过本项目的开发，团队在以下方面获得了显著的技术提升：")
add_body_para("（1）掌握了FastAPI后端框架的使用，理解了依赖注入、中间件、异常处理、API版本管理等企业级开发模式。")
add_body_para("（2）深入理解了AI Agent的设计与实现，包括LLM调用策略、Prompt工程、评分校准、多模式切换、状态持久化等技术。")
add_body_para("（3）掌握了向量检索技术，包括Chroma向量数据库的使用、SentenceTransformer embedding生成、语义匹配评分等。")
add_body_para("（4）理解了RAG知识库的构建方法，包括知识库分层设计、材料标准化、向量化索引、检索结果引用等。")
add_body_para("（5）提升了前端开发能力，包括React Hooks、Recharts图表库、Web Speech API语音交互等。")
add_body_para("（6）学习了软件工程实践，包括Git协作、分支管理、测试驱动开发、持续集成等。")

add_body_para("6.3 不足与展望")
add_body_para("（1）知识库检索尚未深度集成：虽然知识库材料已齐全，但在面试评分和练习计划生成中尚未充分利用知识库中的岗位画像和失败案例，后续需完成Chroma向量化索引和RAG检索管线。")
add_body_para("（2）匹配结果缺少引用来源：岗位匹配结果展示了匹配分数和技能缺口，但未展示匹配依据来自哪些知识库材料。")
add_body_para("（3）TTS语音质量有限：浏览器内置语音合成在中文普通话上表现一般，后续可接入豆包TTS等外部高质量语音服务以提升面试体验。")
add_body_para("（4）数据量有限：当前岗位数据为100条人工生成+24条采集数据，后续可接入真实招聘API实现大规模自动采集。")
add_body_para("（5）多Agent协作未实现：当前面试Agent为单一Agent，未来可引入多Agent辩论、交叉验证等高级模式提升评分准确性。")

add_body_para("6.4 课程学习体会")
add_body_para("本课程「数字媒体技术实训（高级）」通过一个完整的工程项目实践，让我们从需求分析、系统设计、编码实现到测试验收完整地经历了一个软件项目的开发周期。特别是AI Agent、向量检索、RAG知识库等前沿技术的应用实践，极大地拓展了我们的技术视野和工程能力。项目开发过程中，团队协作、版本管理、文档编写等软技能也得到了充分锻炼。")
add_body_para("此外，本项目贴近真实的求职场景，具有较高的实用价值——团队成员在开发过程中也对自身的求职技能有了更清晰的认识，可以说是一次'边做项目边准备求职'的有益实践。")

# ============================================================
# SAVE
# ============================================================
os.makedirs(OUTPUT_DIR, exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"Report saved to: {OUTPUT_PATH}")
print(f"File size: {os.path.getsize(OUTPUT_PATH)} bytes")
