"""
Fill the course report template with ai-career-agent project content.
Reads the .docx template, fills cover info, tables, and body sections.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import copy
from datetime import date

TEMPLATE_PATH = r"C:\Users\17574\Downloads\2026秋-工程实训（高级）-实训报告模板_temp.docx"
OUTPUT_PATH = r"D:\PythonProject\ai-career-agent\docs\2026秋-工程实训（高级）-AI求职Agent实训报告.docx"

doc = Document(TEMPLATE_PATH)

# ============================================================
# Helper: find and replace text across paragraphs
# ============================================================
def replace_text_in_paragraph(para, old, new):
    """Replace text in a paragraph, preserving runs formatting."""
    if old not in para.text:
        return False
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # fallback: text may be split across runs
    full = para.text
    if old in full:
        new_full = full.replace(old, new)
        # Clear all runs and set text in first run
        for i, run in enumerate(para.runs):
            if i == 0:
                run.text = new_full
            else:
                run.text = ""
        return True
    return False


def find_and_replace(doc, old, new):
    """Replace all occurrences of a string in the document."""
    replaced = 0
    # paragraphs
    for para in doc.paragraphs:
        if replace_text_in_paragraph(para, old, new):
            replaced += 1
    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if replace_text_in_paragraph(para, old, new):
                        replaced += 1
    return replaced


def set_cell_text(cell, text, bold=False, size=None, alignment=None):
    """Clear a cell and set its text content."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if alignment is not None:
        para.alignment = alignment
    return run


# ============================================================
# Step 1: Fill cover page info
# ============================================================
print("=== Filling cover page ===")

COVER_REPLACEMENTS = {
    "课   目:":    "课   目:    基于AI Agent的智能求职辅助平台设计与实现",
    "姓   名:":    "姓   名:    何畅",
    "联系电话:":    "联系电话:    待填写",
    "电子邮件:":    "电子邮件:    待填写",
    "2026   年      月      日": f"2026   年   7   月   27   日",
}

for old, new in COVER_REPLACEMENTS.items():
    n = find_and_replace(doc, old, new)
    if n:
        print(f"  Replaced '{old[:20]}...' -> '{new[:40]}...' ({n} occurrences)")

# ============================================================
# Step 2: Fill team member table (Table 1)
# ============================================================
print("\n=== Filling team member table ===")
table1 = doc.tables[0]
members_t1 = [
    ("1", "待填写", "何畅", "数字媒体技术", ""),
    ("2", "待填写", "黄林", "数字媒体技术", ""),
    ("3", "待填写", "牛帅", "数字媒体技术", ""),
    ("4", "待填写", "待填写", "数字媒体技术", ""),
    ("5", "待填写", "待填写", "数字媒体技术", ""),
]
for i, (seq, sid, name, major, score) in enumerate(members_t1):
    row = table1.rows[i + 1]  # skip header
    set_cell_text(row.cells[0], seq, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[1], sid, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[2], name, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[3], major, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[4], score, alignment=WD_ALIGN_PARAGRAPH.CENTER)
print("  Filled 5 member rows")

# ============================================================
# Step 3: Fill contribution table (Table 2)
# ============================================================
print("\n=== Filling contribution table ===")
table2 = doc.tables[1]
members_t2 = [
    ("1", "待填写", "何畅", "项目骨架、认证权限、数据库、后端集成、接口联调、阶段发布", "25%"),
    ("2", "待填写", "黄林", "岗位采集清洗、数据审核、测试数据、失败案例、数据质量文档", "20%"),
    ("3", "待填写", "牛帅", "前端页面、接口联调、交互优化、测试记录、演示材料", "20%"),
    ("4", "待填写", "待填写", "Embedding、向量库、岗位检索、匹配理由、知识库接入", "20%"),
    ("5", "待填写", "待填写", "简历审核、模拟面试、评分反馈、面试题库、报告草稿", "15%"),
]
for i, (seq, sid, name, module, pct) in enumerate(members_t2):
    row = table2.rows[i + 1]
    set_cell_text(row.cells[0], seq, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[1], sid, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[2], name, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[3], module)
    set_cell_text(row.cells[4], pct, alignment=WD_ALIGN_PARAGRAPH.CENTER)
print("  Filled 5 contribution rows")

# ============================================================
# Step 4: Fill body paragraphs (after TOC)
# ============================================================
print("\n=== Filling body content ===")

BODY_SECTIONS = {
    "1、项目概述": """1、项目概述

1.1 项目背景

在当前的就业市场环境下，高校毕业生面临着求职信息分散、简历优化困难、面试准备缺乏针对性等问题。传统的求职方式往往依赖学生个人在各大招聘平台手动搜索岗位、自行修改简历、独自练习面试，整个过程缺乏系统化的指导和智能化的辅助工具。

随着大语言模型（LLM）、向量检索（Vector Search）和检索增强生成（RAG）等人工智能技术的快速发展，构建一个能够自动化处理岗位数据采集与审核、简历智能审核、岗位技能匹配、AI模拟面试和面试报告生成的一体化求职辅助平台成为可能。本项目"基于AI Agent的智能求职辅助平台"正是在这一背景下立项开发。

1.2 项目目标

本项目旨在为求职者（以数字媒体技术及相关专业学生为主要用户群体）提供一个完整的求职辅助工具链，具体目标包括：

（1）岗位数据管理：支持从多个招聘渠道采集岗位数据，进行清洗、去重、技能提取和审核状态流转，构建结构化的岗位数据库。

（2）简历智能审核：支持简历文件上传（PDF/Word/TXT），对简历进行结构化解析，识别优势与不足，给出修改建议。

（3）岗位智能匹配：利用向量检索技术（Chroma + SentenceTransformer），根据简历内容和求职意向，从岗位库中检索最匹配的岗位，输出匹配分数、命中技能、缺失技能和能力缺口分析。

（4）AI模拟面试：构建支持多种面试模式（技术面、HR面、压力面、反馈教练）的AI面试Agent，能够根据候选人简历和岗位要求个性化生成面试题，进行追问、评分和结构化反馈。

（5）面试报告与看板：自动生成面试报告（含维度评分、STAR改写建议、练习计划），并通过数据看板展示求职全过程的统计图表。

1.3 项目意义

本项目的核心价值在于将AI技术实际应用于求职场景，帮助求职者更高效地定位岗位、优化简历、提升面试能力。同时，项目采用了Agent架构、向量检索、RAG知识库等前沿技术，为数字媒体技术专业的学生提供了一个综合性的工程实践平台，有助于理解和掌握现代AI应用系统的设计与开发方法。""",

    "2、需求分析": """2、需求分析

2.1 功能需求

根据课程要求和实际求职场景分析，系统需要满足以下功能需求：

（1）用户认证与权限管理。系统需支持用户注册、登录、JWT Token鉴权，以及基于角色的权限控制（普通学生用户与审核员用户），确保不同角色只能访问对应权限的功能模块。

（2）岗位数据管理。支持岗位数据的JSONL批量导入、单个创建、列表查询、状态流转（pending→approved/rejected），审核员可对岗位进行审核操作。已审核通过的岗位进入向量索引。

（3）简历管理。支持PDF、Word、TXT格式简历文件上传，后端自动解析文本内容。简历支持版本管理，可查看历史版本和当前版本对比。

（4）简历审核。对上传的简历进行规则+LLM双层审核，识别简历中的问题（空泛表达、夸大/绝对化表述、技能缺失等），生成审核报告。

（5）岗位匹配。根据用户简历文本或求职意向，在向量库中检索相似岗位，返回匹配分数、匹配理由、命中技能和缺失技能。

（6）AI面试Agent。支持技术面、HR面、压力面、反馈教练四种面试模式。系统根据候选人简历和目标岗位个性化生成8道面试题，支持追问机制，对回答进行5维度评分，生成STAR改写建议和练习计划。

（7）面试报告。面试结束后自动生成结构化报告，包含总分、维度得分、逐题详情、STAR建议和练习计划。

（8）数据看板。展示简历数、岗位数、面试次数、平均分、热门技能需求、岗位城市分布、最近面试记录等统计图表。

（9）Agent日志。记录Agent每次关键调用的操作类型、状态、耗时、输入输出摘要，用于系统可观测性和调试。

2.2 非功能需求

（1）性能要求：API响应时间应在合理范围内（p95 < 5s），向量检索应在秒级完成。

（2）可靠性要求：LLM离线时系统应能降级运行，通过规则引擎完成基本功能，不影响核心流程。

（3）可用性要求：前端界面应直观易用，面试界面支持语音输入和TTS语音播报。

（4）安全性要求：密码bcrypt加密存储、JWT token鉴权、敏感数据不落地到前端日志。

（5）可扩展性要求：知识库材料（岗位画像、技能词典、面试题库）应独立于代码，支持热更新。

2.3 用户角色分析

系统定义两类用户角色：

（1）学生用户（student角色）：可使用简历上传、简历审核、岗位匹配、模拟面试、面试历史、数据看板等功能。

（2）审核员用户（reviewer角色）：除学生功能外，额外拥有岗位审核、岗位索引重建、系统指标查看等管理权限。""",

    "3、系统设计": """3、系统设计

3.1 系统架构

本系统采用前后端分离的B/S架构，整体分为三层：

（1）前端展示层（React + Vite）：负责用户交互界面，包含登录注册、数据看板、简历管理、岗位管理、岗位匹配、模拟面试、面试历史等9个核心页面。通过RESTful API与后端通信。

（2）后端服务层（Python + FastAPI）：提供业务逻辑处理，包括用户认证、岗位管理、简历审核、向量匹配、面试Agent、报告生成等功能模块。采用依赖注入模式管理数据库会话和用户认证。

（3）数据存储层：使用SQLite作为主数据库（通过SQLAlchemy ORM访问），Chroma作为向量数据库（存储岗位嵌入向量），本地文件系统存储简历上传文件。

3.2 技术架构图

系统核心技术栈如下：

- 后端框架：Python 3.11 + FastAPI 0.110.3
- 数据库：SQLite + SQLAlchemy 2.0.31 ORM
- 向量检索：Chroma + SentenceTransformer（bge-large-zh-v1.5）
- 认证：bcrypt 5.0.0 + python-jose JWT
- LLM集成：OpenAI兼容接口（DeepSeek），支持LLM离线降级
- 前端：React 18 + Vite + Recharts（图表库）
- 简历解析：PyPDF + python-docx
- 语音：浏览器Web Speech API（语音输入）+ Edge TTS / 豆包TTS（语音输出）

3.3 数据库设计

系统设计以下核心数据表：

（1）users表：用户认证信息（username, email, hashed_password, role, is_active, created_at）。

（2）job_postings表：岗位信息（source_id, category, title, company, location, salary_range, education, experience, responsibilities, requirements, skills（JSON）, publish_time, source_site, source_link, collected_at, status, audit_comment）。

（3）resumes表：简历元信息（user_id, title, filename, file_path, source_type, current_version_number, is_default）。

（4）resume_versions表：简历版本内容（resume_id, version_number, content, file_name, created_at）。

（5）resume_audit_reports表：简历审核报告（resume_id, user_id, report_data（JSON）, created_at）。

（6）matching_records表：匹配记录（user_id, job_id, resume_id, total_score, skill_scores, matched_skills, missing_skills, created_at）。

（7）interview_sessions表：面试会话（user_id, job_id, interview_mode, status, score, report_data, started_at, finished_at）。

（8）messages表：聊天消息（session_id, role, content, created_at），其中role=system的消息存储Agent内部状态。

（9）agent_logs表：Agent调用日志（user_id, session_id, operation, status, duration_ms, input_summary, output_summary, error_message, created_at）。

3.4 面试Agent设计

面试Agent是本系统的核心模块，其设计如下：

（1）Agent状态持久化：Agent状态（面试题列表、当前进度、用户回答、评分）序列化为JSON存储在messages表中role=system的记录里，通过_load_agent_state()和_save_agent_state()进行读写。

（2）题目生成策略：采用"简历解析→岗位需求分析→LLM生题"三段式Pipeline。先用LLM解析简历提取技能和项目经历，再分析岗位JD生成面试关注点，最后LLM综合两方信息生成8道个性化题目。题库参考来自interview_question_bank.json中的72道题目（30道通用+42道岗位专属）。

（3）追问决策：采用规则+LLM混合策略。规则层检查回答长度（<50字）、技术关键词覆盖率、量化数据有无、题目要点覆盖度；通过规则判断后由LLM生成个性化追问文本。

（4）评分体系：5维度评分（content_relevance 25分、professional_accuracy 25分、clarity 20分、star_completeness 20分、position_match 10分，满分100分）。LLM评分后经过规则层兜底修正（空泛表达封顶clarity、夸大表达封顶professional_accuracy、漏答要点封顶content_relevance）。

（5）LLM离线降级：所有LLM调用点都有fallback机制，LLM离线时系统可降级为纯规则运行，不影响面试流程完整性。

3.5 知识库设计

系统知识库由四层材料构成：

（1）技能规范层：skill_dictionary.json定义20条技能同义词归一化规则（如Python→[python, Python开发, Python编程]），用于岗位数据清洗时统一技能名称。

（2）岗位能力模型层：role_profiles.json定义6类岗位（前端开发、后端开发、产品经理、运营、算法/机器学习、数字媒体/内容）的能力画像，含must_have技能、preferred技能和evidence_signals。

（3）面试题库层：interview_question_bank.json含30道通用题+42道岗位专属题，覆盖HR面/技术面/压力面/反馈教练四种模式；interview_questions.jsonl含30道带评分点和追问标签的结构化题目。

（4）数据质量层：test_cases.json（10个测试用例）、day2_failure_cases.json（6个失败案例）、job_sources.csv（24条岗位来源追踪）、resume_samples.jsonl（10份脱敏简历样本）、job_jd_samples.jsonl（10条JD样本）。""",

    "4、系统实现": """4、系统实现

4.1 开发环境与工具

- 操作系统：Windows 11
- 开发语言：Python 3.11（后端）、JavaScript/JSX（前端）
- IDE：VS Code / PyCharm
- 版本管理：Git，托管于GitHub
- 包管理：pip（后端）、npm（前端）
- 测试框架：pytest + httpx（后端）
- LLM服务：DeepSeek API（deepseek-chat / deepseek-v4-flash）
- 向量模型：SentenceTransformer bge-large-zh-v1.5

4.2 后端核心实现

4.2.1 FastAPI应用工厂模式

后端采用FastAPI的lifespan模式管理应用生命周期：启动时初始化数据库表、检查向量存储可用性；关闭时清理资源。CORS中间件配置允许前端跨域访问。

4.2.2 异常处理体系

自定义AppException异常类（含status_code、code、message、data、headers），注册四个全局异常处理器（AppException、RequestValidationError、HTTPException、通用Exception），统一返回{code, message, data}格式的API响应。

4.2.3 认证与鉴权

采用bcrypt哈希存储密码（含72字节长度限制检查），JWT Bearer Token鉴权（含type: "access"声明），通过get_current_user()依赖注入获取当前用户，require_roles(*allowed_roles)依赖工厂实现角色权限控制。

4.2.4 向量存储

VectorStore采用线程安全的双重检查锁定单例模式（Double-Checked Locking），使用lazy import避免未安装依赖时导入失败。支持upsert_job()单条索引和search()批量检索，返回余弦相似度分数。

4.2.5 岗位匹配

match_resume_to_jobs()服务函数将简历文本与求职意向拼接为查询字符串，调用VectorStore.search()检索相似岗位，结合技能覆盖度计算综合匹配分数。匹配结果包含命中技能（matched_skills）、缺失技能（missing_skills）和匹配分数。

4.2.6 简历审核

audit_resume_text()采用规则+LLM加权混合审计：规则层检查空泛词（VAGUE_PHRASES）、夸大词（BIASED_PHRASES）、量化数据有无；LLM层进行深度分析。最终分数由30%规则分数+70%LLM分数加权计算。

4.2.7 面试Agent

interview_agent.py为核心实现文件（约1200行），包含以下关键函数：
- start_interview()：解析简历、分析岗位需求、生成8道个性化题目，初始化Agent状态
- evaluate_answer()：接收用户回答，判断是否需要追问或直接评分
- should_followup()：复合规则判断是否需要追问
- _generate_followup()：LLM生成个性化追问
- _score_answer()：LLM评分+规则兜底修正+题目要点覆盖检查
- finish_interview()：计算维度平均分、生成STAR改写建议、练习计划和总评

Agent调用日志通过agent_operation_log()上下文管理器自动记录，包含操作类型、状态、耗时、输入摘要和输出摘要。

4.2.8 知识库集成

系统通过knowledge_base.py模块统一管理知识库访问：role_profile_context()提供岗位画像上下文，get_failure_case_by_scenario()检索失败案例，knowledge_overview()提供知识库统计概览。知识库材料通过scripts/index_knowledge_base.py脚本向量化写入Chroma。

4.3 前端核心实现

4.3.1 页面结构

前端包含9个核心页面：Dashboard（数据看板）、ResumeUpload（简历上传）、ResumeReview（简历审核）、JobCollection（岗位采集）、JobReview（岗位审核）、JobSearchMatch（岗位匹配搜索）、JobComparison（岗位对比）、MockInterview（模拟面试）、InterviewHistory（面试历史）。

4.3.2 Mock API模式

API客户端采用mockFallback模式：后端在线时请求真实API，后端离线时自动降级使用内嵌mock数据返回，保证前端开发不依赖后端状态。

4.3.3 面试交互

模拟面试页面集成了语音输入（浏览器Web Speech API）、摄像头画面采集和TTS语音播报（Edge TTS或豆包TTS），提供沉浸式面试体验。聊天气泡展示每轮得分、追问和反馈。

4.3.4 图表可视化

数据看板使用Recharts库实现6种图表：个人技能分布（水平柱状图）、热门技能需求（垂直柱状图）、个人能力vs岗位要求（雷达图）、多岗位匹配得分（水平柱状图）、面试得分趋势（折线图）、岗位城市分布（饼图）。

4.4 数据工程

4.4.1 岗位数据

系统维护100条中文岗位数据（data/processed/jobs_chinese.jsonl），覆盖6大类岗位、46家公司、11个城市，每条包含19个字段（source_id至audit_comment）。通过scripts/generate_chinese_jobs.py生成，薪资范围按市场行情估算。

4.4.2 数据质量保障

建立10条数据质量测试用例（test_cases.json），覆盖缺字段、缺来源、内容过短、重复检测、相似检测、技能归一化、乱码、日期缺失、审核流转、状态排除等场景。6条系统失败案例（day2_failure_cases.json）用于测试异常输入处理。

4.4.3 岗位审核流程

岗位数据导入后默认status为"pending"，审核员可标记为"approved"（进入向量索引）或"rejected"（排除出匹配），审核意见写入audit_comment字段。""",

    "5、系统测试": """5、系统测试

5.1 测试策略

本系统采用多层次的测试策略：

（1）单元测试：对核心服务函数（面试Agent评分逻辑、规则引擎、岗位清洗、JWT认证等）编写pytest单元测试。

（2）集成测试：对API接口（认证、岗位CRUD、匹配、面试流程等）使用httpx编写集成测试。

（3）前端构建验证：每次提交前执行npm run build确认前端可正常构建。

（4）手动验收测试：对照Day2计划中的验收标准逐项检查功能完整性。

5.2 后端测试

后端测试使用pytest框架，测试文件位于backend/tests/目录。测试覆盖以下核心场景：

（1）认证测试：注册、登录、Token验证、权限控制、密码加密。

（2）岗位管理测试：岗位创建、列表查询、状态审核流转、岗位索引。

（3）岗位匹配测试：匹配接口、匹配历史、技能分数字段。

（4）面试Agent测试：面试启动、回答评估、追问触发、报告生成。

（5）Agent日志测试：日志写入、操作记录、状态记录。

（6）数据看板测试：统计接口、数据正确性。

测试执行命令：cd backend && python -m pytest

所有测试在LLM离线状态下使用mock数据进行，确保测试结果稳定可复现。

5.3 前端测试

前端通过npm run build进行构建验证，确保9个页面组件均能正确编译。各页面使用mockFallback机制，在后端离线时展示模拟数据，验证UI渲染正确性。

5.4 核心功能验收

根据Day2验收标准，逐项验证如下：

（1）完整流程可跑通：注册登录→导入/审核岗位→更新向量索引→简历匹配→Agent面试→结束报告→查看历史。✓

（2）看板至少3个图表模块：个人技能分布、热门技能需求、个人能力vs岗位要求、多岗位匹配得分、面试趋势、城市分布共计6个图表。✓

（3）Agent至少3个工具节点：简历解析、岗位检索、面试出题、追问评分、STAR反馈共计5个节点。✓

（4）面试模式切换：技术面、HR面、压力面、反馈教练共计4种模式。✓

（5）评分结构化：每轮反馈包含得分、优点、问题、改进建议。✓

（6）LLM离线降级：LLM不可用时系统降级为纯规则评分，不影响面试流程完整性。✓

5.5 失败案例测试

系统设计了6个失败测试案例，覆盖以下异常场景：

- FAIL-001（简历太短）：触发should_followup()追问逻辑，预期系统发起追问。
- FAIL-002（岗位无来源）：source_link缺失的岗位，预期审核时不能标记为approved。
- FAIL-003（匹配为空）：技能完全无交集的简历与岗位，预期返回no_match提示。
- FAIL-004（回答过短）：面试回答不足50字，预期触发追问+降低star_completeness分数。
- FAIL-005（含隐私信息）：简历含手机号和邮箱，预期脱敏处理。
- FAIL-006（岗位过期）：超过复核周期的岗位，预期标记为stale。""",

    "6、总结": """6、总结

6.1 项目成果总结

本项目成功构建了一个基于AI Agent的智能求职辅助平台，实现了从岗位数据管理、简历审核、岗位匹配到AI模拟面试的完整闭环。主要成果包括：

（1）后端系统：基于FastAPI构建了完整的RESTful API服务，包含用户认证与鉴权、岗位管理、简历管理、向量匹配、面试Agent、报告生成、数据看板、Agent日志等9大功能模块，代码总量约8000行。

（2）前端系统：基于React+Vite构建了9个交互页面，包含数据看板、简历管理、岗位管理、岗位匹配、模拟面试等核心功能，支持语音输入和TTS语音播报。

（3）数据资产：建立了包含100条中文岗位数据、6类岗位能力画像、20条技能同义词规则、72道面试题、10份脱敏简历样本、10条数据测试用例、6个失败案例的数据资产库。

（4）AI Agent：实现了支持4种面试模式的智能面试Agent，具备题目个性化生成、智能追问、5维度评分、STAR改写建议和练习计划生成能力，且支持LLM离线降级运行。

（5）知识库体系：构建了技能规范、岗位能力模型、面试题库、数据质量标准四层知识库，为RAG检索增强生成奠定基础。

6.2 技术收获

通过本项目的开发，团队在以下方面获得了显著的技术提升：

（1）掌握了FastAPI后端框架的使用，理解了依赖注入、中间件、异常处理、API版本管理等企业级开发模式。

（2）深入理解了AI Agent的设计与实现，包括LLM调用策略、Prompt工程、评分校准、多模式切换、状态持久化等技术。

（3）掌握了向量检索技术，包括Chroma向量数据库的使用、SentenceTransformer embedding生成、语义匹配评分等。

（4）理解了RAG知识库的构建方法，包括知识库分层设计、材料标准化、向量化索引、检索结果引用等。

（5）提升了前端开发能力，包括React Hooks、Recharts图表库、Web Speech API语音交互等。

（6）学习了软件工程实践，包括Git协作、分支管理、测试驱动开发、持续集成等。

6.3 不足与展望

（1）知识库检索尚未深度集成：虽然知识库材料已齐全，但在面试评分和练习计划生成中尚未充分利用知识库中的岗位画像和失败案例，后续需完成Chroma向量化索引和RAG检索管线。

（2）匹配结果缺少引用来源：岗位匹配结果显示匹配分数和技能缺口，但未展示匹配依据来自哪些知识库材料。

（3）TTS语音质量有限：浏览器内置语音合成在中文普通话上表现一般，后续可接入豆包TTS等外部高质量语音服务以提升面试体验。

（4）数据量有限：当前岗位数据为100条人工生成+24条采集数据，后续可接入真实招聘API实现大规模自动采集。

（5）多Agent协作未实现：当前面试Agent为单一Agent，未来可引入多Agent辩论、交叉验证等高级模式提升评分准确性。

6.4 课程学习体会

本课程"数字媒体技术实训（高级）"通过一个完整的工程项目实践，让我们从需求分析、系统设计、编码实现到测试验收完整地经历了一个软件项目的开发周期。特别是AI Agent、向量检索、RAG知识库等前沿技术的应用实践，极大地拓展了我们的技术视野和工程能力。项目开发过程中，团队协作、版本管理、文档编写等软技能也得到了充分锻炼。

此外，本项目贴近真实的求职场景，具有较高的实用价值——团队成员在开发过程中也对自身的求职技能有了更清晰的认识，可以说是一次"边做项目边准备求职"的有益实践。""",
}

# Find the body section start (after TOC)
body_started = False
for para in doc.paragraphs:
    text = para.text.strip()
    if "正文" in text or "1、项目概述" in text or "1��" in text:
        body_started = True
    if body_started:
        for section_title, section_content in BODY_SECTIONS.items():
            if section_title in text or section_title.replace("、", "��")[:3] in text:
                print(f"  Replacing section: {section_title}")
                # Replace this paragraph with section content
                para.text = ""
                para.add_run(section_content)
                break

# Since the old .doc format has mangled encoding, let's rebuild the body section
# Find the paragraph right after TOC heading to insert body content
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    # The body content starts after the TOC
    if any(kw in text for kw in ["正文", "1、项目概述", "1��", "1.项目"]):
        # This is the body section - clear all subsequent body paragraphs
        # and rebuild with our content
        print(f"  Found body section at paragraph {i}: '{text[:50]}'")
        break

# Alternative approach: add body sections as new paragraphs at the end
# since the template's encoding is garbled in the old .doc format
print("\n=== Writing body content at end of document ===")

for title, content in BODY_SECTIONS.items():
    # Split content by double newlines for paragraph separation
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        # Check if it's a heading (starts with a number)
        is_main_heading = (
            line in ("1、项目概述", "2、需求分析", "3、系统设计",
                      "4、系统实现", "5、系统测试", "6、总结")
        )
        is_sub_heading = (
            (line.startswith("1.") or line.startswith("2.") or line.startswith("3.") or
             line.startswith("4.") or line.startswith("5.") or line.startswith("6."))
            and len(line) < 80
        )
        p = doc.add_paragraph()
        if is_main_heading:
            p.style = doc.styles['Heading 1']
            run = p.add_run(line)
            run.font.size = Pt(18)
            run.bold = True
        elif is_sub_heading:
            p.style = doc.styles['Heading 3']
            run = p.add_run(line)
            run.font.size = Pt(13)
            run.bold = True
        else:
            run = p.add_run(line)
            run.font.size = Pt(11)

# ============================================================
# Step 5: Save
# ============================================================
print(f"\n=== Saving to: {OUTPUT_PATH} ===")
doc.save(OUTPUT_PATH)
print("Done! Report saved.")
